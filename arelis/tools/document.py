"""Write a file the user can open — PDF, Word, Excel, CSV, markdown, text.

The model supplies title and body (markdown-ish) or table rows. This tool
renders bytes. It never evals user code. Allow stays (confirm_writes).
Unattended jobs do not get this tool.

Orbit (no room, or a room with no folder) lands under outputs/documents/.
A room with a real project folder lands under that project's documents/.
"""

from __future__ import annotations

import csv
import io
import json
import re
from pathlib import Path
from typing import Any

from arelis.paths import display_path, ensure, outputs_dir
from arelis.rooms import RoomStore
from arelis.tools.base import ToolResult
from arelis.workspace import WorkspaceRoots

_FORMATS = frozenset({"pdf", "docx", "xlsx", "csv", "md", "txt"})
_SOURCE_SUFFIXES = frozenset({".md", ".txt", ".csv", ".markdown"})
_MAX_BODY = 120_000
_MAX_ROWS = 5_000
_MAX_COLS = 40
_SAFE_STEM = re.compile(r"[^a-zA-Z0-9._-]+")
_MD_HEADING = re.compile(r"^(#{1,3})\s+(.*)$")
_MD_BULLET = re.compile(r"^[-*+]\s+(.*)$")
_MD_BOLD = re.compile(r"\*\*(.+?)\*\*")
_TABLE_RULE = re.compile(r"^:?-{3,}:?$")


def _truthy(value: Any) -> bool | None:
    if value is None:
        return None
    if isinstance(value, bool):
        return value
    text = str(value).strip().lower()
    if not text:
        return None
    if text in {"1", "true", "yes", "on", "replace"}:
        return True
    if text in {"0", "false", "no", "off"}:
        return False
    return None


def _unique_dest(directory: Path, stem: str, suffix: str) -> Path:
    candidate = directory / f"{stem}{suffix}"
    if not candidate.exists():
        return candidate
    n = 2
    while True:
        alt = directory / f"{stem}-{n}{suffix}"
        if not alt.exists():
            return alt
        n += 1


def _stem(title: str, filename: str, fmt: str) -> str:
    raw = (filename or "").strip() or (title or "").strip() or "document"
    leaf = Path(raw.replace("\\", "/")).name
    stem = _SAFE_STEM.sub("-", Path(leaf).stem).strip(".-") or "document"
    if Path(leaf).suffix.lower().lstrip(".") == fmt:
        return stem
    return stem


def _plain(text: str) -> str:
    return _MD_BOLD.sub(r"\1", (text or "").replace("\r\n", "\n"))


def _markdown_table(body: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in (body or "").splitlines():
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        if cells and all(_TABLE_RULE.fullmatch(c or "") for c in cells):
            continue
        if cells:
            rows.append(cells[:_MAX_COLS])
        if len(rows) >= _MAX_ROWS:
            break
    return rows


def _table_chunk(lines: list[str]) -> list[list[str]] | None:
    if not lines or not all(ln.lstrip().startswith("|") for ln in lines):
        return None
    rows = _markdown_table("\n".join(lines))
    return rows or None


def _blocks(body: str) -> list[tuple[str, Any]]:
    """Split markdown-ish body into (kind, payload) blocks."""
    text = _plain(body).strip()
    if not text:
        return []
    out: list[tuple[str, Any]] = []
    for chunk in re.split(r"\n{2,}", text):
        lines = [ln.rstrip() for ln in chunk.split("\n") if ln.strip() != ""]
        if not lines:
            continue
        table = _table_chunk(lines)
        if table:
            out.append(("table", table))
            continue
        first = _MD_HEADING.match(lines[0])
        if first and len(lines) == 1:
            level = min(len(first.group(1)), 3)
            out.append((f"h{level}", first.group(2).strip()))
            continue
        bullets = [_MD_BULLET.match(ln) for ln in lines]
        if all(bullets):
            for match in bullets:
                assert match is not None
                out.append(("li", match.group(1).strip()))
            continue
        out.append(("p", "\n".join(lines)))
    return out


def _parse_rows(raw: str, *, body: str) -> list[list[str]]:
    text = (raw or "").strip()
    if text:
        try:
            parsed = json.loads(text)
        except json.JSONDecodeError:
            parsed = None
        if isinstance(parsed, list) and parsed:
            rows: list[list[str]] = []
            for item in parsed[:_MAX_ROWS]:
                if isinstance(item, list):
                    rows.append([str(c) for c in item[:_MAX_COLS]])
                else:
                    rows.append([str(item)])
            return rows
        reader = csv.reader(io.StringIO(text))
        return [row[:_MAX_COLS] for row in reader if any(c.strip() for c in row)][
            :_MAX_ROWS
        ]
    table = _markdown_table(body)
    if table:
        return table
    lines = [ln.strip() for ln in (body or "").splitlines() if ln.strip()]
    if not lines:
        return []
    if any("," in ln or "\t" in ln for ln in lines[:3]):
        reader = csv.reader(io.StringIO("\n".join(lines)))
        return [row[:_MAX_COLS] for row in reader if any(c.strip() for c in row)][
            :_MAX_ROWS
        ]
    return [[ln] for ln in lines[:_MAX_ROWS]]


def _dejavu_paths() -> tuple[str, str]:
    from matplotlib import font_manager

    regular = Path(
        font_manager.findfont(
            font_manager.FontProperties(family="DejaVu Sans", weight="normal")
        )
    )
    bold = Path(
        font_manager.findfont(
            font_manager.FontProperties(family="DejaVu Sans", weight="bold")
        )
    )
    sibling = regular.with_name("DejaVuSans-Bold.ttf")
    if sibling.is_file():
        bold = sibling
    if not regular.is_file():
        raise FileNotFoundError("DejaVu Sans is required to write a PDF.")
    if not bold.is_file():
        bold = regular
    return str(regular), str(bold)


def _pdf_table(pdf: Any, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    padded = [list(row) + [""] * (width - len(row)) for row in rows]
    try:
        with pdf.table(text_align="LEFT") as table:
            for data_row in padded:
                row = table.row()
                for cell in data_row:
                    row.cell(str(cell))
        pdf.ln(2)
        return
    except Exception:
        pass
    usable = pdf.w - pdf.l_margin - pdf.r_margin
    col_w = usable / max(width, 1)
    pdf.set_font("DejaVu", "", 10)
    for data_row in padded:
        y0 = pdf.get_y()
        if y0 > pdf.h - 28:
            pdf.add_page()
            y0 = pdf.get_y()
        x = pdf.l_margin
        bottom = y0
        for cell in data_row:
            pdf.set_xy(x, y0)
            pdf.multi_cell(
                col_w, 5, str(cell), border=1, new_x="RIGHT", new_y="TOP"
            )
            bottom = max(bottom, pdf.get_y())
            x += col_w
        pdf.set_xy(pdf.l_margin, max(bottom, y0 + 5))
    pdf.ln(2)


def _write_pdf(dest: Path, title: str, body: str) -> None:
    from fpdf import FPDF

    regular, bold = _dejavu_paths()
    pdf = FPDF(format="letter", unit="mm")
    pdf.set_auto_page_break(auto=True, margin=22)
    pdf.add_font("DejaVu", fname=regular)
    pdf.add_font("DejaVu", style="B", fname=bold)
    pdf.add_page()
    heading = (title or "").strip() or dest.stem.replace("-", " ")
    pdf.set_font("DejaVu", "B", 18)
    pdf.multi_cell(0, 9, heading, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    blocks = _blocks(body) or [("p", body.strip() or heading)]
    for kind, payload in blocks:
        if kind == "table":
            pdf.set_font("DejaVu", "", 10)
            _pdf_table(pdf, payload)
            continue
        text = str(payload or "")
        if not text:
            continue
        if kind == "h1":
            pdf.set_font("DejaVu", "B", 14)
            pdf.ln(3)
            pdf.multi_cell(0, 8, text, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        elif kind == "h2":
            pdf.set_font("DejaVu", "B", 12)
            pdf.ln(2)
            pdf.multi_cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        elif kind == "h3":
            pdf.set_font("DejaVu", "B", 11)
            pdf.ln(2)
            pdf.multi_cell(0, 6.5, text, new_x="LMARGIN", new_y="NEXT")
        elif kind == "li":
            pdf.set_font("DejaVu", "", 11)
            pdf.multi_cell(0, 6, f"  •  {text}", new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.set_font("DejaVu", "", 11)
            pdf.multi_cell(0, 6, text, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
    pdf.output(str(dest))


def _write_docx(dest: Path, title: str, body: str) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    heading = (title or "").strip() or dest.stem.replace("-", " ")
    doc.add_heading(heading, level=0)
    blocks = _blocks(body) or [("p", body.strip() or heading)]
    for kind, payload in blocks:
        if kind == "table":
            rows = payload
            if not rows:
                continue
            width = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=width)
            table.style = "Table Grid"
            for r_i, row in enumerate(rows):
                for c_i in range(width):
                    table.cell(r_i, c_i).text = str(row[c_i]) if c_i < len(row) else ""
            doc.add_paragraph("")
            continue
        text = str(payload or "")
        if not text:
            continue
        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "h3":
            doc.add_heading(text, level=3)
        elif kind == "li":
            doc.add_paragraph(text, style="List Bullet")
        else:
            para = doc.add_paragraph(text)
            for run in para.runs:
                run.font.size = Pt(11)
    doc.save(str(dest))


def _write_xlsx(dest: Path, rows: list[list[str]], title: str) -> None:
    from openpyxl import Workbook

    book = Workbook()
    sheet = book.active
    sheet.title = (title or "Sheet1")[:31] or "Sheet1"
    for r_i, row in enumerate(rows, start=1):
        for c_i, cell in enumerate(row, start=1):
            sheet.cell(r_i, c_i, cell)
    book.save(str(dest))


def _write_csv(dest: Path, rows: list[list[str]]) -> None:
    with dest.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerows(rows)


def _contained(path: Path, folder: Path) -> bool:
    try:
        path.resolve().relative_to(folder.resolve())
        return True
    except (ValueError, OSError):
        return False


class DocumentTool:
    name = "document"
    description = (
        "Create a file the user can open: PDF, Word (docx), Excel (xlsx), CSV, "
        "markdown, or plain text. In a room with a folder, the file lands in "
        "that project's documents/ directory. Otherwise it lands under "
        "outputs/documents/. Use when they ask to create, make, write, generate, "
        "export, or save a document or spreadsheet. Pass format plus the full "
        "body (markdown is fine). For tables pass rows as JSON arrays or CSV "
        "text. Set replace=true to overwrite the same name (fix / update / "
        "export that). from_path reads an existing .md/.txt/.csv instead of "
        "retyping the body. Do not dump the document into chat. Do not call "
        "doc_extract — that reads an existing PDF. Allow is required."
    )
    risk = "write"
    parameters_schema: dict[str, Any] = {
        "type": "object",
        "properties": {
            "format": {
                "type": "string",
                "enum": ["pdf", "docx", "xlsx", "csv", "md", "txt"],
                "description": "File type to write",
            },
            "title": {
                "type": "string",
                "description": "Document title (also used in the filename)",
            },
            "body": {
                "type": "string",
                "description": "Full document text. Markdown headings, bullets, and tables are fine.",
            },
            "rows": {
                "type": "string",
                "description": "Spreadsheet data: JSON list of lists, or CSV text",
            },
            "filename": {
                "type": "string",
                "description": "Output file name only (no folders)",
            },
            "replace": {
                "type": "string",
                "description": "true to overwrite the same name; false to write a new file beside it",
            },
            "from_path": {
                "type": "string",
                "description": "Existing .md/.txt/.csv to render instead of body",
            },
        },
        "required": ["format"],
    }

    def __init__(
        self,
        workspace: WorkspaceRoots | None = None,
        rooms: RoomStore | None = None,
    ) -> None:
        self.workspace = workspace
        self.rooms = rooms

    def drop_dir(self) -> Path:
        return ensure(outputs_dir() / "documents")

    def room_docs_dir(self) -> Path | None:
        """Project/documents when a room with a live folder is open."""
        room = None if self.rooms is None else self.rooms.active
        if room is None or not room.root or self.workspace is None:
            return None
        entry = self.workspace.root_named(room.root)
        if entry is None:
            return None
        return ensure(entry.path / "documents")

    def out_dir(self) -> tuple[Path, str]:
        """(folder, short where-it-landed note)."""
        room_dir = self.room_docs_dir()
        if room_dir is not None:
            return room_dir, "this room's documents folder"
        room = None if self.rooms is None else self.rooms.active
        if room is not None and room.root:
            return (
                self.drop_dir(),
                "the shared drop tray — this room's folder is not a project any more",
            )
        if room is not None:
            return (
                self.drop_dir(),
                "the shared drop tray — this room has no folder",
            )
        return self.drop_dir(), "the shared drop tray (outputs/documents)"

    def preview_path(self, args: dict[str, Any] | None = None) -> Path:
        """Where this call will write, without writing it."""
        args = args or {}
        fmt = str(args.get("format") or "pdf").strip().lower().lstrip(".")
        aliases = {
            "word": "docx",
            "doc": "docx",
            "excel": "xlsx",
            "spreadsheet": "xlsx",
            "markdown": "md",
            "text": "txt",
        }
        fmt = aliases.get(fmt, fmt) or "pdf"
        if fmt not in _FORMATS:
            fmt = "pdf"
        folder, _where = self.out_dir()
        stem = _stem(
            str(args.get("title") or ""),
            str(args.get("filename") or ""),
            fmt,
        )
        replace = _truthy(args.get("replace")) is True
        dest = folder / f"{stem}.{fmt}"
        if dest.exists() and not replace:
            dest = _unique_dest(folder, stem, f".{fmt}")
        return dest

    def _read_source(self, raw: str) -> str:
        text = (raw or "").strip()
        if not text:
            raise ValueError("from_path is empty.")
        path: Path | None = None
        if self.workspace is not None:
            try:
                path = self.workspace.resolve_read(text).path
            except (ValueError, PermissionError, OSError):
                path = None
        candidate = Path(text)
        if path is None and candidate.is_file():
            resolved = candidate.resolve()
            folder, _where = self.out_dir()
            if _contained(resolved, folder) or _contained(resolved, self.drop_dir()):
                path = resolved
        if path is None or not path.is_file():
            raise ValueError(f"Cannot read {text!r} as a source file.")
        if path.suffix.lower() not in _SOURCE_SUFFIXES:
            raise ValueError("from_path must be a markdown, text, or CSV file.")
        body = path.read_text(encoding="utf-8", errors="replace")
        if len(body) > _MAX_BODY:
            raise ValueError(f"source file is too long (max {_MAX_BODY} characters).")
        return body

    async def run(self, **kwargs: Any) -> ToolResult:
        fmt = str(kwargs.get("format") or "").strip().lower().lstrip(".")
        aliases = {
            "word": "docx",
            "doc": "docx",
            "excel": "xlsx",
            "spreadsheet": "xlsx",
            "markdown": "md",
            "text": "txt",
        }
        fmt = aliases.get(fmt, fmt)
        if fmt not in _FORMATS:
            return ToolResult(
                ok=False,
                output="format must be pdf, docx, xlsx, csv, md, or txt.",
            )
        title = str(kwargs.get("title") or "").strip()
        body = str(kwargs.get("body") or "")
        rows_raw = str(kwargs.get("rows") or "")
        from_path = str(kwargs.get("from_path") or "").strip()
        if from_path:
            try:
                body = self._read_source(from_path)
            except Exception as exc:
                return ToolResult(ok=False, output=str(exc))
            if not title:
                title = Path(from_path.replace("\\", "/")).stem.replace("-", " ")
        if len(body) > _MAX_BODY:
            return ToolResult(
                ok=False,
                output=f"body is too long (max {_MAX_BODY} characters).",
            )
        if fmt in {"csv", "xlsx"}:
            try:
                rows = _parse_rows(rows_raw, body=body)
            except Exception as exc:
                return ToolResult(ok=False, output=f"Could not read table data: {exc}")
            if not rows:
                return ToolResult(
                    ok=False,
                    output="A spreadsheet needs rows (JSON list of lists or CSV) or a markdown table in body.",
                )
        else:
            rows = []
            if not body.strip() and not title:
                return ToolResult(
                    ok=False,
                    output="A document needs title or body — nothing to write.",
                )
            if not body.strip():
                body = title

        folder, where = self.out_dir()
        stem = _stem(title, str(kwargs.get("filename") or ""), fmt)
        replace = _truthy(kwargs.get("replace")) is True
        dest = folder / f"{stem}.{fmt}"
        replaced = dest.exists() and replace
        if dest.exists() and not replace:
            dest = _unique_dest(folder, stem, f".{fmt}")
        if not _contained(dest, folder):
            return ToolResult(
                ok=False,
                output="That filename cannot leave the documents folder.",
            )
        try:
            if fmt == "pdf":
                _write_pdf(dest, title, body)
            elif fmt == "docx":
                _write_docx(dest, title, body)
            elif fmt == "xlsx":
                _write_xlsx(dest, rows, title)
            elif fmt == "csv":
                _write_csv(dest, rows)
            elif fmt == "md":
                heading = (
                    f"# {title}\n\n"
                    if title and not body.lstrip().startswith("#")
                    else ""
                )
                dest.write_text(heading + body.strip() + "\n", encoding="utf-8")
            else:
                dest.write_text(body.strip() + "\n", encoding="utf-8")
        except Exception as exc:
            return ToolResult(ok=False, output=f"Could not write {fmt}: {exc}")

        shown = display_path(dest)
        pages = ""
        if fmt == "pdf":
            try:
                from pypdf import PdfReader

                n = len(PdfReader(str(dest)).pages)
                pages = f" ({n} page{'s' if n != 1 else ''})"
            except Exception:
                pages = ""
        verb = "Replaced" if replaced else "Wrote"
        return ToolResult(
            ok=True,
            output=(
                f"{verb} {shown}{pages} in {where}. "
                "Open that file — do not paste the document into chat."
            ),
            data={
                "path": shown,
                "abs_path": str(dest.resolve()),
                "format": fmt,
                "title": title,
                "replaced": replaced,
                "where": where,
            },
        )
